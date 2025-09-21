import json
import pandas as pd
from rich.console import Console
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def save_json(report_data, file_path="baseline_report.json"):
    """Saves the report data to a JSON file."""
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(report_data, f, indent=4)
    console = Console()
    console.print(f"[green][INFO][/green] JSON report saved to {file_path}")

def save_csv(report_data, file_path="baseline_report.csv"):
    """Saves the report data to a CSV file."""
    all_features = set(report_data["baseline_features"]) | set(report_data["non_baseline_features"])
    data = []
    for feat in sorted(list(all_features)):
        status = "Baseline" if feat in report_data["baseline_features"] else "Non-Baseline"
        data.append({"Feature": feat, "Status": status})
    
    df = pd.DataFrame(data)
    df.to_csv(file_path, index=False)
    console = Console()
    console.print(f"[green][INFO][/green] CSV report saved to {file_path}")

def save_word(report_data, file_path="baseline_report.docx"):
    """Saves the report data to a Word document."""
    doc = Document()
    doc.add_heading('Baseline Compatibility Report', 0)
    doc.add_paragraph(f"Total files scanned: {report_data['total_files_scanned']}")

    table = doc.add_table(rows=1, cols=2, style='LightShading-Accent1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Status'

    all_features = sorted(list(set(report_data["baseline_features"]) | set(report_data["non_baseline_features"])))
    for feature in all_features:
        row_cells = table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = "Baseline" if feature in report_data["baseline_features"] else "Non-Baseline"
    
    doc.save(file_path)
    console = Console()
    console.print(f"[green][INFO][/green] Word report saved to {file_path}")

def save_pdf(report_data, file_path="baseline_report.pdf"):
    """Saves the report data to a PDF file."""
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    y = height - 40
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(30, y, "Baseline Compatibility Report")
    y -= 20
    
    c.setFont("Helvetica", 12)
    c.drawString(30, y, f"Total files scanned: {report_data['total_files_scanned']}")
    y -= 30

    c.setFont("Helvetica-Bold", 12)
    c.drawString(30, y, "Features:")
    y -= 15
    
    c.setFont("Helvetica", 10)
    all_features = sorted(list(set(report_data["baseline_features"]) | set(report_data["non_baseline_features"])))

    for feature in all_features:
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 40

        status = "Baseline" if feature in report_data["baseline_features"] else "Non-Baseline"
        c.drawString(40, y, f"- {feature}: {status}")
        y -= 12
    # Prepare report data
   